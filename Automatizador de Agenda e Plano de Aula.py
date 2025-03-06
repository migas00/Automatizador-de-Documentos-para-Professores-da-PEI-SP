from openpyxl import load_workbook
from docx import Document

# Here it will start loading the Files and Create the Documents

wb = load_workbook('2. Anos Finais - Escopo-sequência 2025.xlsx') # here it will load the Escopo file
ws = wb['Língua Inglesa'] # here it will load the subject sheet

# A few inputs

date_doc = input('Qual a data do plano de aula? (modelo: ddmm a ddmm) ') # here it will ask for the date of the plan
num_doc = int(input('Qual o numero do plano de aula? (modelo: 1, 2, 3, 4, etc.) ')) # here it will ask for the number of the plan in order to find it in the Escopo file

def creating_doc(grade, date, num): # this function makes the entire document
    new_doc = Document(f'PLANO QUINZENAL MODELO.docx') # here it will load the template
    edit_date = new_doc.paragraphs[3] # here it will edit the date of the plan in the document
    edit_date.text = f'PLANO DE AULA QUINZENAL 2025 - {int(date[:2]):02d}/{int(date[3:5]):02d} a {int(date[6:8]):02d}/{int(date[9:11]):02d}'
    table0 = new_doc.tables[0] # here it will edit the grade of the plan in the document
    table0.cell(0,0).text = f'Ano/Série: {grade}º Ano' # here it will edit the grade of the Plano de Aula in the document
    if grade == 7: # it checks if the grade is the 7th
        first_class_title = ws.cell(2 * (num_doc-1) + 15, 8).value # it will get the title of the first class of the 7th grade
        second_class_title = ws.cell(2 * (num_doc-1) + 16, 8).value # it will get the title of the second class of the 7th grade    
        first_class_content = ws.cell(2 * (num_doc-1) + 15, 9).value # it will get the content of the first class of the 7th grade  
        second_class_content = ws.cell(2 * (num_doc-1) + 16, 9).value # it will get the content of the second class of the 7th grade    
        first_class_goals = ws.cell(2 * (num_doc-1) + 15, 10).value # it will get the goals of the first class of the 7th grade 
        second_class_goals = ws.cell(2 * (num_doc-1) + 16, 10).value # it will get the goals of the second class of the 7th grade   
    elif grade == 8: # it checks if the grade is the 8th
        first_class_title = ws.cell(2 * (num_doc-1) + 35, 8).value # it will get the title of the first class of the 8th grade
        second_class_title = ws.cell(2 * (num_doc-1) + 36, 8).value # it will get the title of the second class of the 8th grade        
        first_class_content = ws.cell(2 * (num_doc-1) + 35, 9).value # it will get the content of the first class of the 8th grade  
        second_class_content = ws.cell(2 * (num_doc-1) + 36, 9).value # it will get the content of the second class of the 8th grade    
        first_class_goals = ws.cell(2 * (num_doc-1) + 35, 10).value # it will get the goals of the first class of the 8th grade 
        second_class_goals = ws.cell(2 * (num_doc-1) + 36, 10).value # it will get the goals of the second class of the 8th grade   
    elif grade == 9: # it checks if the grade is the 9th
        first_class_title = ws.cell(2 * (num_doc-1) + 55, 8).value # it will get the title of the first class of the 9th grade
        second_class_title = ws.cell(2 * (num_doc-1) + 56, 8).value # it will get the title of the second class of the 9th grade            
        first_class_content = ws.cell(2 * (num_doc-1) + 55, 9).value # it will get the content of the first class of the 9th grade    
        second_class_content = ws.cell(2 * (num_doc-1) + 56, 9).value # it will get the content of the second class of the 9th grade    
        first_class_goals = ws.cell(2 * (num_doc-1) + 55, 10).value # it will get the goals of the first class of the 9th grade      
        second_class_goals = ws.cell(2 * (num_doc-1) + 56, 10).value # it will get the goals of the second class of the 9th grade      
    if first_class_title == 'Paltaforma EF': # check if the title is miswritten and then corrects it
        first_class_title = 'Plataforma EF' 
    elif second_class_title == 'Paltaforma EF':
        second_class_title = 'Plataforma EF'
    table1 = new_doc.tables[1] # it will start editing the table
    if input(f'Deseja personalizar a metodologia ou avaliação do {grade}º ano? (s/n) ') == 's': # asks if the user wants to personalize the methodology or evaluation of the plan
        if input(f'Deseja personalizar a metodologia da PRIMEIRA semana do {grade}º ano? (s/n) ') == 's': # asks if the user wants to personalize the methodology of the first week of the plan
            table1.cell(1, 3).text = input('Digite a metodologia: ')
        if input(f'Deseja personalizar a avaliação da PRIMEIRA semana do {grade}º ano? (s/n) ') == 's': # asks if the user wants to personalize the evaluation of the first week of the plan
            table1.cell(1, 4).text = input('Digite a avaliação: ')  
        if input(f'Deseja personalizar a metodologia da SEGUNDA semana do {grade}º ano? (s/n) ') == 's': # asks if the user wants to personalize the methodology of the second week of the plan
            table1.cell(2, 3).text = input('Digite a metodologia: ')
        if input(f'Deseja personalizar a avaliação da SEGUNDA semana do {grade}º ano? (s/n) ') == 's': # asks if the user wants to personalize the evaluation of the second week of the plan
            table1.cell(2, 4).text = input('Digite a avaliação: ')
    else:
        if first_class_title == 'Plataforma EF' and second_class_title == 'Plataforma EF': # it checks if its an Plataforma EF class, which changes the methodology and evaluation of the plan
            table1.cell(1, 3).text = 'Sala de aula invertida \nOs alunos conduzem as aulas e o professor, monitora-os em suas dificuldades.'
            table1.cell(1, 4).text = 'Acompanhamento da evolução do aluno durante as aulas na plataforma.'
            table1.cell(2, 3).text = 'Sala de aula invertida \nOs alunos conduzem as aulas e o professor, monitora-os em suas dificuldades.'
            table1.cell(2, 4).text = 'Acompanhamento da evolução do aluno durante as aulas na plataforma.'
        elif first_class_title == 'Plataforma EF': # same thing, but checks if only the first class is Plataforma EF
            table1.cell(1, 3).text = 'Sala de aula invertida \nOs alunos conduzem as aulas e o professor, monitora-os em suas dificuldades.'
            table1.cell(1, 4).text = 'Acompanhamento da evolução do aluno durante as aulas na plataforma.'
            table1.cell(2, 3).text = 'Aula Explicativa e Expositiva'
            table1.cell(2, 4).text = ' Atividade no caderno do aluno.\nCorreção da atividade e participação no processo.'
        elif second_class_title == 'Plataforma EF': # same thing, but checks if only the second class is Plataforma EF
            table1.cell(1, 3).text = 'Aula Explicativa e Expositiva'
            table1.cell(1, 4).text = ' Atividade no caderno do aluno.\nCorreção da atividade e participação no processo.'
            table1.cell(2, 3).text = 'Sala de aula invertida \nOs alunos conduzem as aulas e o professor, monitora-os em suas dificuldades.'
            table1.cell(2, 4).text = 'Acompanhamento da evolução do aluno durante as aulas na plataforma.'
        else: # if it's not Plataforma EF, it will just save the document with the given name
            table1.cell(1, 3).text = 'Aula Explicativa e Expositiva'
            table1.cell(1, 4).text = 'Atividade no caderno do aluno.\nCorreção da atividade e participação no processo.'
            table1.cell(2, 3).text = 'Aula Explicativa e Expositiva'
            table1.cell(2, 4).text = 'Atividade no caderno do aluno.\nCorreção da atividade e participação no processo.'
    table1.cell(1, 0).text = first_class_title # edits the first class title
    table1.cell(2, 0).text = second_class_title # edits the second class title
    table1.cell(1, 1).text = first_class_content # edits the first class content
    table1.cell(2, 1).text = second_class_content # edits the second class content
    table1.cell(1, 2).text = first_class_goals # edits the first class goals
    table1.cell(2, 2).text = second_class_goals # edits the second class goals
    new_doc.save(f'Plano de Aula - {grade}º Ano - Inglês - {date} - {num}.docx') # saves the document with the given name

creating_doc(7, date_doc, num_doc) # creates the document for the 7th grade
creating_doc(8, date_doc, num_doc) # creates the document for the 8th grade
creating_doc(9, date_doc, num_doc) # creates the document for the 9th grade



# ////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// #



# here it will create the sheets for the Agenda



numberstomonths: dict = {'01' : 'Janeiro', '02' : 'Fevereiro', '03' : 'Março', '04' : 
                        'Abril', '05' : 'Maio', '06' : 'Junho', '07' : 'Julho', 
                        '08' : 'Agosto', '09' : 'Setembro', '10' : 'Outubro', 
                        '11' : 'Novembro', '12' : 'Dezembro'}

# a few inputs

st_date_1st = int(input('Defina o dia e mês que começa a planilha da primeira semana: (modelo: ddmm) ')) 
end_date_1st = int(input('Defina o dia e mês que termina a planilha da primeira semana: (modelo: ddmm) '))
st_date_2nd = int(input('Defina o dia e mês que começa a planilha da segunda semana: (modelo: ddmm) '))
end_date_2nd = int(input('Defina o dia e mês que termina a planilha da segunda semana: (modelo: ddmm) '))

# transforms the inputs getting only the day

starting_day_sheet1 = st_date_1st // 100
ending_day_sheet1 = end_date_1st // 100
starting_day_sheet2 = st_date_2nd // 100
ending_day_sheet2 = end_date_2nd // 100

# transforms the inputs getting only the month

starting_month_sheet1 = st_date_1st%100
ending_month_sheet1 = end_date_1st%100
starting_month_sheet2 = st_date_2nd%100
ending_month_sheet2 = end_date_2nd%100


def build_doc(date_doc, serie, row):  # gets the title of classes from the docx file
    doc = Document(f'Plano de Aula - {serie}º Ano - Inglês - {date_doc} - {num_doc}.docx') # loads the Plano de Aula 
    class_title = doc.tables[1].cell(row, 0).text # gets the title of the class
    return class_title

title7_sheet1 = build_doc(date_doc, 7, 1) # gets the title of the first week's classes
title8_sheet1 = build_doc(date_doc, 8, 1) # gets the title of the first week's classes
title9_sheet1 = build_doc(date_doc, 9, 1) # gets the title of the first week's classes

title7_sheet2 = build_doc(date_doc, 7, 2) # gets the title of the second week's classes
title8_sheet2 = build_doc(date_doc, 8, 2) # gets the title of the second week's classes
title9_sheet2 = build_doc(date_doc, 9, 2) # gets the title of the second week's classes

# starting to edit the sheet

wb = load_workbook('NOME_DO(A)_PROFESSOR(A)_AGENDA.xlsx') # loads the Agenda
def create_sheets(wb, sheet_name):
    template = wb['1702 a 2102'] # makes the first sheet as the template for all new sheets (the first sheet is where the teacher edits his classes and their schedule based on the day,
    # for example, if he has a class with the 7th grade at 9:45 A.M., he will edit the first sheet at 9:45 A.M. and the new sheets will have the same schedule)
    wb.copy_worksheet(template).title = sheet_name # creates a new sheet with the inputed name and the same schedule as the template

create_sheets(wb, f'{st_date_1st:04d} a {end_date_1st:04d}') # creates the first week's sheet
create_sheets(wb, f'{st_date_2nd:04d} a {end_date_2nd:04d}') # creates the second week's sheet
worksheet1 = wb[f'{st_date_1st:04d} a {end_date_1st:04d}'] # loads the first week's sheet to edit it
worksheet2 = wb[f'{st_date_2nd:04d} a {end_date_2nd:04d}'] # loads the second week's sheet to edit it

# the teacher has 2 types of classes: normal and platform. The first one is at the classroom and the second one is at the library

normal_classes7 = ['A22', 'I15', 'M29'] # these are the cells of the normal classes of the 7th grade
normal_classes8 = ['A43', 'A64', 'I50', 'Q15', 'Q64'] # these are the cells of the normal classes of the 8th grade
normal_classes9 = ['A15', 'E15', 'E29', 'I29', 'M50' ] # these are the cells of the normal classes of the 9th grade

for i in normal_classes7:
    worksheet1[i].value = title7_sheet1
    worksheet2[i].value = title7_sheet2
for i in normal_classes8:
    worksheet1[i].value = title8_sheet1
    worksheet2[i].value = title8_sheet2    
for i in normal_classes9:
    worksheet1[i].value = title9_sheet1
    worksheet2[i].value = title9_sheet2
platform_classes = ['A29', 'I29', 'M43', 'A50', 'A71', 'I64', 'Q22', 'Q71', 'Q43', 'E22', 'E43', 'I43', 'M64'] # these are the cells of the platform classes
for i in platform_classes:
    worksheet1[i].value = 'Trilha de Aprendizagem Individual'
    worksheet2[i].value = 'Trilha de Aprendizagem Individual'

# time to edit the respective days in the sheet

def edit_days(ws, starting_day, ending_day, starting_month, ending_month): # this function edits the days in which the sheet will work
    starting_day = f'{starting_day:02d}'
    ending_day = f'{ending_day:02d}'
    starting_month = f'{starting_month:02d}'
    ending_month = f'{ending_month:02d}'
    ws['K9'].value = f'Mês: {numberstomonths[starting_month]}'
    ws['A12'].value = f'{starting_day}/{starting_month}'
    ws['Q12'].value = f'{ending_day}/{ending_month}'
    if starting_month != ending_month: # if the starting and ending months are different, it means that some classes will be in the inputed month and others will be in the following month
        rest = 4 - int(ending_day)
        if rest == 0: # if the last day of the sheet is 4, then only the first day is on the inputed month and the others are on the following month
            ws['E12'].value = f'01/{ending_month}'
            ws['I12'].value = f'02/{ending_month}'
            ws['M12'].value = f'03/{ending_month}'
        elif rest == 1: # if the last day of the sheet is 3, then the first and second days are on the inputed month and the others are on the following month
            ws['E12'].value = f'{int(starting_day) + 1:02d}/{starting_month}'
            ws['I12'].value = f'01/{ending_month}'
            ws['M12'].value = f'02/{ending_month}'
        elif rest == 2: # if the last day of the sheet is 2, then the first, second and third days are on the inputed month and the others are on the following month
            ws['E12'].value = f'{int(starting_day) + 1:02d}/{starting_month}'
            ws['I12'].value = f'{int(starting_day) + 2:02d}/{starting_month}'
            ws['M12'].value = f'01/{ending_month}'
        elif rest == 3: # if the last day of the sheet is 1, then only one day is on the following month
            ws['E12'].value = f'{int(starting_day) + 1:02d}/{starting_month}'
            ws['I12'].value = f'{int(starting_day) + 2:02d}/{starting_month}'
            ws['M12'].value = f'{int(starting_day) + 3:02d}/{starting_month}'
    else: # if the starting and ending months are the same, then all days are on the inputed month and nothing different is done.
        ws['E12'].value = f'{int(starting_day) + 1:02d}/{starting_month}'
        ws['I12'].value = f'{int(starting_day) + 2:02d}/{starting_month}'
        ws['M12'].value = f'{int(starting_day) + 3:02d}/{starting_month}'

edit_days(worksheet1, starting_day_sheet1, ending_day_sheet1, starting_month_sheet1, ending_month_sheet1) # edits the first week's sheet
edit_days(worksheet2, starting_day_sheet2, ending_day_sheet2, starting_month_sheet2, ending_month_sheet2) # edits the second week's sheet

wb.save('NOME_DO(A)_PROFESSOR(A)_AGENDA.xlsx') # saves the Agenda